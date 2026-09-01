import io
import logging
from io import BytesIO
from typing import Optional

import pdfplumber
import fitz  # PyMuPDF
from docx import Document
import PyPDF2

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

import os

MAX_EXTRACTED_CHARS = 100000  # 100k chars limit
MAX_PDF_PAGES = 50

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using a resilient multi-tier fallback pipeline:
    1. pdfplumber (highest layout quality)
    2. PyMuPDF / fitz (fast fallback)
    3. PyPDF2 (legacy fallback)
    """
    text = ""
    # Tier 1: pdfplumber
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            if len(pdf.pages) > MAX_PDF_PAGES:
                raise ValueError(f"PDF exceeds maximum page limit of {MAX_PDF_PAGES} pages.")
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                if len(text) > MAX_EXTRACTED_CHARS:
                    break
    except Exception as e:
        if isinstance(e, ValueError) and "maximum page limit" in str(e):
            raise e
        logger.warning(f"pdfplumber extraction failed: {e}. Trying PyMuPDF fallback.")
        text = ""

    # Tier 2: PyMuPDF (fitz)
    if not text.strip():
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            if len(doc) > MAX_PDF_PAGES:
                raise ValueError(f"PDF exceeds maximum page limit of {MAX_PDF_PAGES} pages.")
            for page in doc:
                text += page.get_text() + "\n"
                if len(text) > MAX_EXTRACTED_CHARS:
                    break
        except Exception as e:
            if isinstance(e, ValueError) and "maximum page limit" in str(e):
                raise e
            logger.warning(f"PyMuPDF extraction failed: {e}. Trying PyPDF2 fallback.")
            text = ""

    # Tier 3: PyPDF2
    if not text.strip():
        try:
            pdf_file = BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file)
            if len(reader.pages) > MAX_PDF_PAGES:
                raise ValueError(f"PDF exceeds maximum page limit of {MAX_PDF_PAGES} pages.")
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                if len(text) > MAX_EXTRACTED_CHARS:
                    break
        except Exception as e:
            if isinstance(e, ValueError) and "maximum page limit" in str(e):
                raise e
            logger.error(f"PyPDF2 extraction failed: {e}")
            text = ""

    if not text.strip():
        raise ValueError("Could not extract legible text from PDF file. The file may be image-only or corrupt.")

    return text[:MAX_EXTRACTED_CHARS].strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX paragraph nodes and table cells."""
    try:
        doc = Document(BytesIO(file_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))
        full_text = "\n".join(text_parts)
        if not full_text.strip():
            raise ValueError("DOCX document contains no text.")
        return full_text[:MAX_EXTRACTED_CHARS]
    except Exception as e:
        logger.error(f"python-docx extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX file: {e}")

def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain TXT file with encoding fallback."""
    for encoding in ['utf-8', 'latin-1', 'cp1252']:
        try:
            text = file_bytes.decode(encoding)
            if text.strip():
                return text[:MAX_EXTRACTED_CHARS].strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("Failed to decode text file. Ensure it is encoded in UTF-8 or ASCII.")

def parse_resume_file(file_bytes: bytes, filename: str) -> str:
    """
    Main entry point for extracting text from resume files (PDF, DOCX, TXT).
    Sanitizes filenames against path traversal, validates magic bytes and file size.
    """
    if not file_bytes:
        raise ValueError("Uploaded file is empty.")
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File size exceeds maximum limit of {MAX_FILE_SIZE_BYTES // (1024*1024)}MB.")

    # Path traversal protection
    safe_filename = os.path.basename(filename.replace("\\", "/"))
    fn_lower = safe_filename.lower()

    # Magic byte validation
    if fn_lower.endswith(".pdf"):
        if not file_bytes.startswith(b"%PDF-"):
            raise ValueError("Invalid file content for .pdf extension (magic bytes mismatch).")
        return extract_text_from_pdf(file_bytes)
    elif fn_lower.endswith(".docx") or fn_lower.endswith(".doc"):
        if not file_bytes.startswith(b"PK\x03\x04") and not fn_lower.endswith(".doc"):
            raise ValueError("Invalid file content for .docx extension (magic bytes mismatch).")
        return extract_text_from_docx(file_bytes)
    elif fn_lower.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file format for '{safe_filename}'. Supported formats: .pdf, .docx, .txt")

